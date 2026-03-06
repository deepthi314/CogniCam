import io
from datetime import datetime
from typing import Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_cam_report(data: Dict[str, Any]) -> bytes:
    """
    Generate Credit Appraisal Memorandum (CAM) Word document.
    
    Args:
        data: Dictionary containing all appraisal data
        
    Returns:
        Word document as bytes
    """
    print("📄 Starting CAM report generation...")
    
    # Create document with A4 page setup
    doc = Document()
    section = doc.sections[0]
    section.page_height = Inches(11.7)  # A4 height
    section.page_width = Inches(8.3)     # A4 width
    section.left_margin = Inches(1.0)    # 2.5cm margins
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    # HEADER TABLE
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(4.0)
    header_table.columns[1].width = Inches(4.0)
    
    # Left cell - COGNICAM branding
    left_cell = header_table.cell(0, 0)
    left_cell.paragraphs[0].clear()
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    _set_cell_bg(left_cell, "#1A1F3A")
    
    cognicam_run = left_para.add_run("⚡ COGNICAM")
    cognicam_run.font.size = Pt(16)
    cognicam_run.font.bold = True
    cognicam_run.font.color.rgb = RGBColor(255, 255, 255)
    
    subtitle_run = left_para.add_run("\nContext-Aware Credit Appraisal Engine")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(200, 200, 200)
    
    # Right cell - Hackathon branding
    right_cell = header_table.cell(0, 1)
    right_cell.paragraphs[0].clear()
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    _set_cell_bg(right_cell, "#4F46E5")
    
    hackathon_run = right_para.add_run("IIT Hyderabad Hackathon 2025")
    hackathon_run.font.size = Pt(14)
    hackathon_run.font.bold = True
    hackathon_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # TITLE SECTION
    doc.add_paragraph().add_run().add_break()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_para.add_run("CREDIT APPRAISAL MEMORANDUM (CAM)")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 31, 58)
    
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_name = data.get("company_name", "Unknown Company")
    today = datetime.now().strftime("%d %B %Y")
    subtitle_run = subtitle_para.add_run(f"{company_name} | Date: {today}")
    subtitle_run.font.size = Pt(11)
    
    # SECTION 1 — EXECUTIVE SUMMARY
    _add_section_header(doc, "EXECUTIVE SUMMARY")
    
    # 4-column decision table
    decision_table = doc.add_table(rows=1, cols=4)
    decision_table.autofit = False
    
    # Set column widths
    for i in range(4):
        decision_table.columns[i].width = Inches(2.0)
    
    # Get recommendation data
    recommendation = data.get("recommendation", {})
    decision = recommendation.get("decision", "PENDING")
    final_score = recommendation.get("final_score", 0)
    risk_grade = recommendation.get("risk_grade", "N/A")
    confidence = recommendation.get("confidence", 0) * 100
    
    # Decision cell
    decision_cell = decision_table.cell(0, 0)
    decision_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell_bg(decision_cell, _get_decision_color(decision))
    decision_run = decision_cell.paragraphs[0].add_run(f"DECISION\n{decision}")
    decision_run.font.size = Pt(14)
    decision_run.font.bold = True
    decision_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Score cell
    score_cell = decision_table.cell(0, 1)
    score_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell_bg(score_cell, "#4F46E5")
    score_run = score_cell.paragraphs[0].add_run(f"CREDIT SCORE\n{final_score:.1f}/100")
    score_run.font.size = Pt(14)
    score_run.font.bold = True
    score_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Grade cell
    grade_cell = decision_table.cell(0, 2)
    grade_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell_bg(grade_cell, _get_grade_color(risk_grade))
    grade_run = grade_cell.paragraphs[0].add_run(f"RISK GRADE\n{risk_grade}")
    grade_run.font.size = Pt(14)
    grade_run.font.bold = True
    grade_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Confidence cell
    conf_cell = decision_table.cell(0, 3)
    conf_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    _set_cell_bg(conf_cell, "#10B981")
    conf_run = conf_cell.paragraphs[0].add_run(f"CONFIDENCE\n{confidence:.1f}%")
    conf_run.font.size = Pt(14)
    conf_run.font.bold = True
    conf_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Borrower details
    financial_data = data.get("financial_data", {})
    borrower_para = doc.add_paragraph()
    borrower_para.add_run("Borrower Details: ").bold = True
    borrower_para.add_run(f"{company_name} | GSTIN: {financial_data.get('gstin', 'N/A')} | FY: {financial_data.get('financial_year', 'N/A')}")
    
    # Credit terms
    credit_limit_cr = recommendation.get("credit_limit_cr", 0)
    interest_rate = recommendation.get("interest_rate", 0)
    tenor = recommendation.get("tenor", "N/A")
    
    terms_para = doc.add_paragraph()
    terms_para.add_run("Recommended Credit Terms: ").bold = True
    if decision != "DECLINE":
        terms_para.add_run(f"Limit: ₹{credit_limit_cr:.2f} Cr @ {interest_rate}% for {tenor}")
    else:
        terms_para.add_run("Application Declined")
    
    # SECTION 2 — FIVE Cs SCORECARD
    _add_section_header(doc, "FIVE Cs SCORECARD")
    
    scores_table = doc.add_table(rows=6, cols=4)
    scores_table.autofit = False
    scores_table.columns[0].width = Inches(1.5)
    scores_table.columns[1].width = Inches(1.0)
    scores_table.columns[2].width = Inches(1.0)
    scores_table.columns[3].width = Inches(3.0)
    
    # Header row
    headers = ["C-Parameter", "Score", "Grade", "Key Rationale"]
    for i, header in enumerate(headers):
        cell = scores_table.cell(0, i)
        _set_cell_bg(cell, "#1A1F3A")
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = cell.paragraphs[0].add_run(header)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    scores = data.get("scores", {})
    rationales = data.get("rationales", {})
    c_params = ["character", "capacity", "capital", "collateral", "conditions"]
    
    for i, param in enumerate(c_params, 1):
        score = scores.get(param, 0)
        rationale = rationales.get(param, "N/A")
        grade_letter, grade_text, grade_color = _grade_from_score(score)
        
        # Parameter name
        scores_table.cell(i, 0).paragraphs[0].add_run(param.title()).bold = True
        
        # Score
        scores_table.cell(i, 1).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        scores_table.cell(i, 1).paragraphs[0].add_run(str(score))
        
        # Grade
        grade_cell = scores_table.cell(i, 2)
        _set_cell_bg(grade_cell, grade_color)
        grade_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        grade_run = grade_cell.paragraphs[0].add_run(f"{grade_letter}\n{grade_text}")
        grade_run.font.bold = True
        grade_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Rationale (truncated)
        rationale_text = rationale[:120] + "..." if len(rationale) > 120 else rationale
        scores_table.cell(i, 3).paragraphs[0].add_run(rationale_text)
    
    # SECTION 3 — GST COMPLIANCE ANALYSIS
    _add_section_header(doc, "GST COMPLIANCE ANALYSIS")
    
    gstr_flags = data.get("gstr_flags", {})
    risk_level = gstr_flags.get("risk_level", "Low")
    flags = gstr_flags.get("flags", [])
    
    gst_para = doc.add_paragraph()
    gst_para.add_run("Risk Level: ").bold = True
    gst_para.add_run(f"{risk_level}")
    
    if flags:
        for flag in flags:
            flag_para = doc.add_paragraph()
            flag_para.add_run("⚠ ").bold = True
            flag_para.add_run(flag).bold = True
    else:
        doc.add_paragraph("✓ No GSTR anomalies detected")
    
    # SECTION 4 — LEGAL & NEWS INTELLIGENCE
    _add_section_header(doc, "LEGAL & NEWS INTELLIGENCE")
    
    # 4.1 Litigation Cases
    litigation_heading = doc.add_paragraph()
    litigation_heading.add_run("4.1 Litigation Cases").bold = True
    
    litigation_cases = data.get("litigation_cases", [])
    if litigation_cases:
        cases_table = doc.add_table(rows=len(litigation_cases) + 1, cols=4)
        cases_table.autofit = False
        
        # Header
        case_headers = ["Case No.", "Court", "Status", "Nature"]
        for i, header in enumerate(case_headers):
            cell = cases_table.cell(0, i)
            _set_cell_bg(cell, "#1A1F3A")
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = cell.paragraphs[0].add_run(header)
            header_run.font.bold = True
            header_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        for i, case in enumerate(litigation_cases, 1):
            cases_table.cell(i, 0).paragraphs[0].add_run(case.get("case_number", "N/A"))
            cases_table.cell(i, 1).paragraphs[0].add_run(case.get("court", "N/A"))
            
            status_cell = cases_table.cell(i, 2)
            if case.get("status", "").lower() == "pending":
                _set_cell_bg(status_cell, "#FEF3C7")  # Light yellow
            status_cell.paragraphs[0].add_run(case.get("status", "N/A"))
            
            cases_table.cell(i, 3).paragraphs[0].add_run(case.get("nature", "N/A"))
    else:
        doc.add_paragraph("✓ No litigation cases found")
    
    # 4.2 News Sentiment
    news_heading = doc.add_paragraph()
    news_heading.add_run("4.2 News Sentiment Analysis").bold = True
    
    news_articles = data.get("news_articles", [])
    for article in news_articles[:5]:  # Limit to 5 articles
        sentiment = article.get("sentiment", "NEUTRAL")
        confidence = article.get("confidence", 0) * 100
        headline = article.get("headline", "No headline")
        
        news_para = doc.add_paragraph()
        if sentiment == "POSITIVE":
            news_para.add_run("[+] ").bold = True
        elif sentiment == "NEGATIVE":
            news_para.add_run("[-] ").bold = True
        else:
            news_para.add_run("[~] ").bold = True
        
        news_para.add_run(f"SENTIMENT: {headline} ({confidence:.0f}%)")
        
        if article.get("credit_risk_flag", False):
            risk_para = doc.add_paragraph()
            risk_para.add_run("⚠ RISK FLAG DETECTED").bold = True
    
    # SECTION 5 — AI EXPLAINABILITY
    _add_section_header(doc, "AI EXPLAINABILITY")
    
    shap_data = data.get("shap_explanation", {})
    if shap_data:
        shap_table = doc.add_table(rows=6, cols=3)
        shap_table.autofit = False
        
        # Header
        shap_headers = ["Factor", "SHAP Value", "Impact Direction"]
        for i, header in enumerate(shap_headers):
            cell = shap_table.cell(0, i)
            _set_cell_bg(cell, "#1A1F3A")
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            header_run = cell.paragraphs[0].add_run(header)
            header_run.font.bold = True
            header_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Data rows
        shap_values = shap_data.get("shap_values", {})
        factors = ["character", "capacity", "capital", "collateral", "conditions"]
        
        for i, factor in enumerate(factors, 1):
            value = shap_values.get(factor, 0)
            
            shap_table.cell(i, 0).paragraphs[0].add_run(factor.title()).bold = True
            shap_table.cell(i, 1).paragraphs[0].add_run(f"{value:+.3f}")
            
            impact_cell = shap_table.cell(i, 2)
            if value > 0.01:
                _set_cell_bg(impact_cell, "#D1FAE5")  # Light green
                impact_cell.paragraphs[0].add_run("Positive")
            elif value < -0.01:
                _set_cell_bg(impact_cell, "#FEE2E2")  # Light red
                impact_cell.paragraphs[0].add_run("Negative")
            else:
                impact_cell.paragraphs[0].add_run("Neutral")
        
        # Explanation text
        explanation = shap_data.get("explanation_text", "")
        if explanation:
            exp_para = doc.add_paragraph()
            exp_para.add_run(explanation)
    
    # SECTION 6 — CREDIT RECOMMENDATION
    _add_section_header(doc, "CREDIT RECOMMENDATION")
    
    rec_table = doc.add_table(rows=1, cols=2)
    rec_table.autofit = False
    rec_table.columns[0].width = Inches(4.0)
    rec_table.columns[1].width = Inches(4.0)
    
    # Credit limit cell
    limit_cell = rec_table.cell(0, 0)
    _set_cell_bg(limit_cell, "#1A1F3A")
    limit_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    limit_para = limit_cell.paragraphs[0]
    limit_para.add_run("RECOMMENDED CREDIT LIMIT\n").font.size = Pt(12)
    limit_para.add_run().font.color.rgb = RGBColor(255, 255, 255)
    
    if decision != "DECLINE":
        limit_para.add_run(f"₹{credit_limit_cr:.2f} Cr").font.size = Pt(16).bold = True
    else:
        limit_para.add_run("DECLINED").font.size = Pt(16).bold = True
    
    # Rate & tenor cell
    rate_cell = rec_table.cell(0, 1)
    _set_cell_bg(rate_cell, _get_decision_color(decision))
    rate_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    rate_para = rate_cell.paragraphs[0]
    rate_para.add_run("INTEREST RATE & TENOR\n").font.size = Pt(12)
    rate_para.add_run().font.color.rgb = RGBColor(255, 255, 255)
    
    if decision != "DECLINE":
        rate_para.add_run(f"{interest_rate}% for {tenor}").font.size = Pt(16).bold = True
    else:
        rate_para.add_run("N/A").font.size = Pt(16).bold = True
    
    # SECTION 7 — FIELD OFFICER NOTES
    _add_section_header(doc, "FIELD OFFICER NOTES")
    
    field_note = data.get("field_note", "")
    if field_note:
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(field_note)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(100, 100, 100)
    else:
        doc.add_paragraph("No field officer notes provided.")
    
    # FOOTER
    doc.add_paragraph().add_run().add_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_run = footer_para.add_run(f"Generated by COGNICAM AI Engine | {today} | Confidential — For Bank Use Only")
    footer_run.font.size = Pt(8)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    
    # Save to bytes
    print("📝 Finalizing CAM document...")
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    
    print(f"✅ CAM report generated: {len(doc_bytes.getvalue())} bytes")
    return doc_bytes.read()

def _set_cell_bg(cell, hex_color):
    """Set cell background color using OxmlElement"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), hex_color.lstrip('#'))
    cell._tc.get_or_add_tcPr().append(shading)

def _add_section_header(doc, title):
    """Add section header with indigo bottom border"""
    para = doc.add_paragraph()
    para.add_run(title).font.size = Pt(13).bold = True
    para.runs[0].font.color.rgb = RGBColor(26, 31, 58)
    
    # Add bottom border
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pBdr.set(qn('w:bottom'), qn('w:double'))
    pBdr.set(qn('w:bottom-color'), '#4F46E5')
    pBdr.set(qn('w:bottom-size'), '6')
    pBdr.set(qn('w:bottom-space'), '1')
    pPr.append(pBdr)

def _grade_from_score(score):
    """Get grade letter, text, and color from score"""
    if score >= 75:
        return "A", "Excellent", "#10B981"
    elif score >= 50:
        return "B", "Satisfactory", "#F59E0B"
    else:
        return "C", "Poor", "#EF4444"

def _get_decision_color(decision):
    """Get color based on decision"""
    if decision == "APPROVE":
        return "#10B981"  # Emerald
    elif "CONDITIONAL" in decision:
        return "#F59E0B"  # Amber
    else:
        return "#EF4444"  # Red

def _get_grade_color(grade):
    """Get color based on grade"""
    if grade == "A":
        return "#10B981"  # Emerald
    elif grade == "B":
        return "#F59E0B"  # Amber
    else:
        return "#EF4444"  # Red
